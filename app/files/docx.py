"""JSON structuré → octets .docx (python-docx)."""

from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PETROLE = RGBColor(0x0B, 0x2D, 0x2A)
GRIS = RGBColor(0x55, 0x6B, 0x66)


def _set_cell_background(cell, hex_color: str) -> None:
    """Ombrage de cellule (clear/auto — jamais un noir plein)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_bottom_border(paragraph) -> None:
    """Filet horizontal via bordure de paragraphe."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B2D2A")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_document(payload: dict, sources: list[dict]) -> bytes:
    """Transforme le formulaire JSON + sources en octets Word."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    head = doc.add_paragraph()
    r = head.add_run("DYNEFF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = PETROLE

    sub = doc.add_paragraph()
    rs = sub.add_run("Direction des Ressources Humaines")
    rs.font.size = Pt(10)
    rs.font.color.rgb = GRIS
    _add_bottom_border(sub)

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_run(date.today().strftime("Le %d/%m/%Y")).font.size = Pt(10)

    if payload.get("destinataire"):
        doc.add_paragraph().add_run(str(payload["destinataire"])).font.size = Pt(11)

    if payload.get("objet"):
        obj = doc.add_paragraph()
        lab = obj.add_run("Objet : ")
        lab.bold = True
        obj.add_run(str(payload["objet"]))

    if payload.get("type") == "note" and payload.get("titre"):
        doc.add_heading(str(payload["titre"]), level=1)

    for bloc in payload.get("blocs", []) or []:
        t = bloc.get("type")
        if t == "titre":
            doc.add_heading(str(bloc.get("texte", "")), level=2)
        elif t == "paragraphe":
            doc.add_paragraph(str(bloc.get("texte", "")))
        elif t == "liste":
            for item in bloc.get("items", []) or []:
                doc.add_paragraph(str(item), style="List Bullet")
        elif t == "tableau":
            entetes = bloc.get("entetes", []) or []
            lignes = bloc.get("lignes", []) or []
            if entetes:
                table = doc.add_table(rows=1, cols=len(entetes))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = table.rows[0].cells
                for i, texte in enumerate(entetes):
                    hdr[i].text = ""
                    run = hdr[i].paragraphs[0].add_run(str(texte))
                    run.bold = True
                    _set_cell_background(hdr[i], "E8ECEF")
                for ligne in lignes:
                    cells = table.add_row().cells
                    for i, val in enumerate(ligne):
                        if i < len(cells):
                            cells[i].text = str(val)

    if payload.get("signature"):
        doc.add_paragraph()
        doc.add_paragraph().add_run(str(payload["signature"])).font.size = Pt(11)

    if sources:
        doc.add_paragraph()
        rule = doc.add_paragraph()
        _add_bottom_border(rule)
        sh = doc.add_paragraph().add_run("Sources")
        sh.bold = True
        sh.font.size = Pt(10)
        for i, s in enumerate(sources, start=1):
            libelle = " · ".join(
                str(x)
                for x in [
                    s.get("document"),
                    s.get("section"),
                    f"p.{s.get('page')}" if s.get("page") else None,
                ]
                if x
            )
            doc.add_paragraph().add_run(f"[{i}] {libelle}").font.size = Pt(9)

    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = "Document généré par l'Assistant RH Dyneff — à relire avant signature."
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x90, 0x96)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
